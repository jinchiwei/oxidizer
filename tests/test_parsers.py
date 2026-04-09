"""Tests for oxidizer document parsers."""
import io
import os
import tempfile

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from oxidizer.parsers.markdown_parser import Section, _classify_context, parse_markdown
from oxidizer.parsers.latex_parser import parse_latex
from oxidizer.parsers.docx_parser import parse_docx


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_docx(paragraphs: list[tuple[str, str]]) -> str:
    """Create a temporary .docx file from a list of (style_name, text) tuples.

    Returns the path to the temporary file. Caller is responsible for cleanup.
    """
    doc = Document()
    for style_name, text in paragraphs:
        if style_name == "Normal":
            doc.add_paragraph(text)
        else:
            doc.add_heading(text, level=int(style_name.split()[-1]))
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


# ---------------------------------------------------------------------------
# _classify_context
# ---------------------------------------------------------------------------

class TestClassifyContext:
    def test_introduction(self):
        assert _classify_context("Introduction") == "intro"

    def test_background(self):
        assert _classify_context("Background and Motivation") == "intro"

    def test_methods(self):
        assert _classify_context("Methods") == "methods"

    def test_materials_and_methods(self):
        assert _classify_context("Materials and Methods") == "methods"

    def test_experimental_setup(self):
        assert _classify_context("Experimental Setup") == "methods"

    def test_results(self):
        assert _classify_context("Results") == "results"

    def test_discussion(self):
        assert _classify_context("Discussion") == "discussion"

    def test_conclusion(self):
        assert _classify_context("Conclusion") == "discussion"

    def test_other(self):
        assert _classify_context("Appendix A") == "other"

    def test_case_insensitive(self):
        assert _classify_context("INTRODUCTION") == "intro"


# ---------------------------------------------------------------------------
# parse_markdown
# ---------------------------------------------------------------------------

class TestParseMarkdown:
    def test_single_heading(self):
        text = "# Introduction\n\nThis is the intro body."
        sections = parse_markdown(text)
        assert len(sections) == 1
        s = sections[0]
        assert s.heading == "Introduction"
        assert s.body == "This is the intro body."
        assert s.context == "intro"
        assert s.level == 1

    def test_multiple_headings(self):
        text = (
            "# Introduction\n\nIntro body.\n\n"
            "## Methods\n\nMethods body.\n\n"
            "### Results\n\nResults body."
        )
        sections = parse_markdown(text)
        assert len(sections) == 3
        assert sections[0].heading == "Introduction"
        assert sections[0].level == 1
        assert sections[1].heading == "Methods"
        assert sections[1].level == 2
        assert sections[1].context == "methods"
        assert sections[2].heading == "Results"
        assert sections[2].level == 3
        assert sections[2].context == "results"

    def test_no_headings_returns_single_section(self):
        text = "Just some plain text with no headings at all."
        sections = parse_markdown(text)
        assert len(sections) == 1
        s = sections[0]
        assert s.heading == ""
        assert s.level == 0
        assert s.context == "other"
        assert s.body == text.strip()

    def test_body_stripped(self):
        text = "# Discussion\n\n   Some text.   \n\n"
        sections = parse_markdown(text)
        assert sections[0].body == "Some text."

    def test_context_mapping_discussion(self):
        text = "# Conclusions\n\nWrapping up."
        sections = parse_markdown(text)
        assert sections[0].context == "discussion"

    def test_deep_heading_level(self):
        text = "###### Deep Section\n\nDeep body."
        sections = parse_markdown(text)
        assert sections[0].level == 6

    def test_empty_string_no_headings(self):
        sections = parse_markdown("")
        assert len(sections) == 1
        assert sections[0].body == ""

    def test_section_dataclass_fields(self):
        s = Section(heading="Foo", body="Bar", context="other", level=2)
        assert s.heading == "Foo"
        assert s.body == "Bar"
        assert s.context == "other"
        assert s.level == 2


# ---------------------------------------------------------------------------
# parse_latex
# ---------------------------------------------------------------------------

class TestParseLaTeX:
    def test_single_section(self):
        text = r"\section{Introduction}" + "\n\nIntro body."
        sections = parse_latex(text)
        assert len(sections) == 1
        assert sections[0].heading == "Introduction"
        assert sections[0].context == "intro"
        assert sections[0].level == 1
        assert sections[0].body == "Intro body."

    def test_multiple_sections(self):
        text = (
            r"\section{Introduction}" + "\n\nIntro body.\n\n"
            r"\section{Methods}" + "\n\nMethods body.\n\n"
            r"\subsection{Experimental Setup}" + "\n\nSetup body."
        )
        sections = parse_latex(text)
        assert len(sections) == 3
        assert sections[0].heading == "Introduction"
        assert sections[0].level == 1
        assert sections[1].heading == "Methods"
        assert sections[1].level == 1
        assert sections[1].context == "methods"
        assert sections[2].heading == "Experimental Setup"
        assert sections[2].level == 2
        assert sections[2].context == "methods"

    def test_subsubsection_level(self):
        text = r"\subsubsection{Detail}" + "\n\nDetail text."
        sections = parse_latex(text)
        assert sections[0].level == 3

    def test_no_sections_fallback(self):
        text = r"Just some \LaTeX{} preamble text without sections."
        sections = parse_latex(text)
        assert len(sections) == 1
        assert sections[0].heading == ""
        assert sections[0].level == 0
        assert sections[0].context == "other"

    def test_context_results(self):
        text = r"\section{Results}" + "\n\nOur findings."
        sections = parse_latex(text)
        assert sections[0].context == "results"

    def test_context_discussion(self):
        text = r"\section{Discussion}" + "\n\nWe discuss."
        sections = parse_latex(text)
        assert sections[0].context == "discussion"

    def test_body_stripped(self):
        text = r"\section{Methods}" + "\n\n   Body text.   \n\n"
        sections = parse_latex(text)
        assert sections[0].body == "Body text."


# ---------------------------------------------------------------------------
# parse_docx
# ---------------------------------------------------------------------------

class TestParseDocx:
    def test_happy_path_with_headings(self):
        path = _make_docx([
            ("Heading 1", "Introduction"),
            ("Normal", "This is the introduction body."),
            ("Heading 1", "Methods"),
            ("Normal", "These are the methods."),
        ])
        try:
            sections = parse_docx(path)
            assert len(sections) == 2
            assert sections[0].heading == "Introduction"
            assert sections[0].context == "intro"
            assert sections[0].level == 1
            assert "introduction body" in sections[0].body

            assert sections[1].heading == "Methods"
            assert sections[1].context == "methods"
            assert sections[1].level == 1
            assert "methods" in sections[1].body
        finally:
            os.unlink(path)

    def test_no_headings_returns_single_section(self):
        path = _make_docx([
            ("Normal", "Paragraph one."),
            ("Normal", "Paragraph two."),
        ])
        try:
            sections = parse_docx(path)
            assert len(sections) == 1
            assert sections[0].heading == ""
            assert sections[0].level == 0
            assert sections[0].context == "other"
            assert "Paragraph one" in sections[0].body
        finally:
            os.unlink(path)

    def test_heading_levels(self):
        path = _make_docx([
            ("Heading 1", "Results"),
            ("Normal", "Top-level result."),
            ("Heading 2", "Sub-results"),
            ("Normal", "Sub-level result."),
        ])
        try:
            sections = parse_docx(path)
            assert sections[0].level == 1
            assert sections[1].level == 2
            assert sections[0].context == "results"
        finally:
            os.unlink(path)

    def test_discussion_context(self):
        path = _make_docx([
            ("Heading 1", "Conclusions"),
            ("Normal", "We conclude that..."),
        ])
        try:
            sections = parse_docx(path)
            assert sections[0].context == "discussion"
        finally:
            os.unlink(path)

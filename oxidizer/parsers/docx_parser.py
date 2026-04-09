"""Parse .docx files into sections using python-docx."""
from docx import Document
from oxidizer.parsers.markdown_parser import Section, _classify_context

# Heading style name prefixes used by python-docx
_HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5", "Heading 6"}


def _heading_level(style_name: str) -> int:
    """Return the numeric level for a Word heading style (1-6), or 0 if not a heading."""
    for i in range(1, 7):
        if style_name == f"Heading {i}":
            return i
    return 0


def parse_docx(path: str) -> list[Section]:
    """Parse a .docx file into a list of Sections split on heading styles.

    If no headings are found the entire document text is returned as a single
    Section with an empty heading and context "other".
    """
    doc = Document(path)

    # Collect (level, heading_text, body_paragraphs) groups
    groups: list[tuple[int, str, list[str]]] = []

    for para in doc.paragraphs:
        level = _heading_level(para.style.name)
        if level:
            groups.append((level, para.text.strip(), []))
        else:
            if groups:
                groups[-1][2].append(para.text)
            # paragraphs before the first heading are collected below

    # Handle paragraphs that appear before any heading
    pre_heading: list[str] = []
    for para in doc.paragraphs:
        level = _heading_level(para.style.name)
        if level:
            break
        pre_heading.append(para.text)

    sections: list[Section] = []

    if not groups:
        # No headings — return everything as one section
        body = "\n".join(p.text for p in doc.paragraphs).strip()
        return [Section(heading="", body=body, context="other", level=0)]

    for level, heading, body_paras in groups:
        body = "\n".join(body_paras).strip()
        context = _classify_context(heading)
        sections.append(Section(heading=heading, body=body, context=context, level=level))

    return sections
